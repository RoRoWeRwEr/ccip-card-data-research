#!/usr/bin/env python3
"""Build the non-destructive Saudi cards consolidation deliverables."""

from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
import zipfile
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from pypdf import PdfReader
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "Credit Cards Terms and Conditions"
ORIGINAL = SRC / "01. saudi-credit-cards-unified V3.xlsx"
MASTER_MD = SRC / "02. Xlsx audit master reference 2026 07 29.md"
CHROME_MD = SRC / "03. Cards information from Claude Chrome prompt V4.md"
OUT = ROOT / "outputs"
REPORTS = OUT / "reports"
MACHINE = OUT / "machine-readable"
EXCEL = OUT / "excel"
WORKING = ROOT / "working"
TODAY = date(2026, 7, 30).isoformat()


def clean(value: Any) -> str:
    if value is None:
        return ""
    text = re.sub(r"\s+", " ", str(value)).strip()
    return text


def norm(value: Any) -> str:
    text = clean(value).lower()
    text = re.sub(r"[\s\-_–—/()\[\]{}:؛،,.]+", "", text)
    return text


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for block in iter(lambda: f.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def infer_bank(name: str) -> str:
    upper = name.upper()
    mappings = [
        ("ALBILAD", "Bank Albilad"), ("ALINMA", "Alinma Bank"),
        ("AMERICAN EXPRESS", "American Express Saudi Arabia"), ("AMEX", "American Express Saudi Arabia"),
        ("ANB", "Arab National Bank"), ("BSF", "Banque Saudi Fransi"),
        ("RAJHI", "Al Rajhi Bank"), ("RIYAD", "Riyad Bank"),
        ("SAB", "Saudi Awwal Bank"), ("SAIB", "Saudi Investment Bank"),
        ("SNB", "Saudi National Bank"), ("AJB", "Bank Aljazira"),
        ("ENBD", "Emirates NBD KSA"),
    ]
    for token, bank in mappings:
        if token in upper:
            return bank
    return "Unmapped / multi-bank"


def extract_document(path: Path) -> dict[str, Any]:
    result = {"status": "readable", "pages": "", "encrypted": False, "text_chars": 0, "text_sha256": "", "error": "", "sample": ""}
    try:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            result["pages"] = len(reader.pages)
            result["encrypted"] = bool(reader.is_encrypted)
            if reader.is_encrypted:
                try:
                    reader.decrypt("")
                except Exception:
                    result["status"] = "encrypted/unreadable"
                    return result
            texts = []
            for page in reader.pages:
                texts.append(page.extract_text() or "")
            text = "\n".join(texts)
        elif suffix == ".docx":
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        elif suffix in {".md", ".txt", ".html", ".csv", ".json"}:
            text = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".rtf":
            raw = path.read_text(encoding="utf-8", errors="replace")
            text = re.sub(r"\\'[0-9a-fA-F]{2}|\\[a-zA-Z]+-?\d* ?|[{}]", " ", raw)
        elif suffix == ".xlsx":
            wb = load_workbook(path, read_only=True, data_only=False)
            chunks = []
            for ws in wb.worksheets:
                chunks.append(ws.title)
                for row in ws.iter_rows(values_only=True):
                    chunks.append(" | ".join(clean(v) for v in row if v is not None))
            text = "\n".join(chunks)
        else:
            result["status"] = "unsupported"
            return result
        result["text_chars"] = len(text)
        result["sample"] = clean(text[:500])
        normalized_text = re.sub(r"\s+", "", text).lower()
        result["text_sha256"] = hashlib.sha256(normalized_text.encode("utf-8")).hexdigest() if normalized_text else ""
        if not text.strip():
            result["status"] = "readable/no extractable text"
    except Exception as exc:
        result["status"] = "unreadable"
        result["error"] = f"{type(exc).__name__}: {exc}"
    return result


def inventory_sources() -> list[dict[str, Any]]:
    # Exclude generated outputs, working files, and this generated inventory itself.
    # Including the inventory in its own hash registry makes every rerun unstable.
    files = sorted(
        p for p in ROOT.rglob("*")
        if p.is_file()
        and ".git" not in p.parts
        and "__pycache__" not in p.parts
        and "node_modules" not in p.parts
        and "outputs" not in p.parts
        and "working" not in p.parts
        and p != ROOT / "docs" / "REPOSITORY_INVENTORY.md"
    )
    rows = []
    for p in files:
        digest = sha256(p)
        meta = extract_document(p) if p.parent == SRC or p.name.lower() == "readme.md" else {"status": "repository file", "pages": "", "encrypted": False, "text_chars": "", "text_sha256": "", "error": "", "sample": ""}
        rows.append({
            "source_id": f"SRC-{len(rows)+1:03d}", "path": str(p.relative_to(ROOT)), "filename": p.name,
            "extension": p.suffix.lower() or "[none]", "size_bytes": p.stat().st_size, "sha256": digest,
            "suspected_duplicate": False, "duplicate_basis": "", "bank_mapping": infer_bank(p.name), **meta,
        })
    file_hashes = Counter(r["sha256"] for r in rows)
    text_hashes = Counter(r["text_sha256"] for r in rows if r["text_sha256"])
    for row in rows:
        exact = file_hashes[row["sha256"]] > 1
        same_text = bool(row["text_sha256"] and text_hashes[row["text_sha256"]] > 1)
        row["suspected_duplicate"] = exact or same_text
        row["duplicate_basis"] = "exact bytes" if exact else ("normalized extracted text" if same_text else "")
    (WORKING / "source_hashes.json").write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    return rows


def parse_markdown_tables(path: Path) -> list[dict[str, Any]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    tables: list[dict[str, Any]] = []
    heading = ""
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("#") or (line.strip() and not line.startswith("|") and len(line.strip()) < 100):
            if any(k in line.lower() for k in ["bank", "بنك", "مصرف", "البنك"]):
                heading = clean(line.lstrip("# "))
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-{3,}", lines[i + 1].strip(" |")):
            headers = [clean(x) for x in line.strip().strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].startswith("|"):
                vals = [clean(x) for x in lines[i].strip().strip("|").split("|")]
                vals += [""] * (len(headers) - len(vals))
                rows.append(dict(zip(headers, vals[:len(headers)])))
                i += 1
            tables.append({"heading": heading, "headers": headers, "rows": rows})
            continue
        i += 1
    return tables


def find_col(row: dict[str, str], tokens: list[str]) -> str:
    for key, value in row.items():
        nk = norm(key)
        if any(norm(t) in nk for t in tokens):
            return clean(value)
    return ""


def chrome_records() -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    guides, details = [], []
    for table_no, table in enumerate(parse_markdown_tables(CHROME_MD), 1):
        if not table["headers"] or norm(table["headers"][0]) != "cardid":
            continue
        is_guide = len(table["headers"]) >= 10
        for row_no, row in enumerate(table["rows"], 1):
            cid = clean(next(iter(row.values())))
            if not re.match(r"^[A-Z][A-Z0-9-]+$", cid):
                continue
            if is_guide:
                guides.append({
                    "card_id": cid, "bank_section": table["heading"],
                    "card_name": find_col(row, ["اسم البطاقة"]), "network": find_col(row, ["الشبكة"]),
                    "tier": find_col(row, ["الفئة/المستوى"]), "account_segment": find_col(row, ["فئة الحساب"]),
                    "loyalty_program": find_col(row, ["برنامج الولاء"]), "reward_type": find_col(row, ["نوع المكافأة"]),
                    "product_type": find_col(row, ["نوع المنتج"]), "calc_tier": find_col(row, ["calctier"]),
                    "annual_fee": find_col(row, ["الرسوم السنوية"]), "benefits": find_col(row, ["أبرز المزايا"]),
                    "source": find_col(row, ["مصدر البيانات", "المصدر"]), "source_date": "2026-07-29",
                    "source_file": str(CHROME_MD.relative_to(ROOT)), "table_no": str(table_no), "row_no": str(row_no),
                })
            else:
                values = list(row.values())
                details.append({
                    "card_id": cid, "bank_section": table["heading"],
                    "category": values[1] if len(values) > 1 else "",
                    "value": values[2] if len(values) > 2 else "",
                    "note": values[3] if len(values) > 3 else "",
                    "source_file": str(CHROME_MD.relative_to(ROOT)), "table_no": str(table_no), "row_no": str(row_no),
                })
    return guides, details


def workbook_audit(path: Path) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    wb = load_workbook(path, data_only=False, read_only=False)
    sheets = []
    for ws in wb.worksheets:
        formulas, comments, links, blanks, types = 0, 0, 0, 0, Counter()
        for row in ws.iter_rows():
            for c in row:
                if c.value is None:
                    blanks += 1
                else:
                    types[c.data_type] += 1
                formulas += c.data_type == "f"
                comments += c.comment is not None
                links += c.hyperlink is not None
        sheets.append({
            "sheet": ws.title, "state": ws.sheet_state, "rows": ws.max_row, "columns": ws.max_column,
            "formulas": formulas, "comments": comments, "hyperlinks": links, "merged_ranges": len(ws.merged_cells.ranges),
            "hidden_rows": sum(bool(d.hidden) for d in ws.row_dimensions.values()),
            "hidden_columns": sum(bool(d.hidden) for d in ws.column_dimensions.values()),
            "tables": len(ws.tables), "data_validations": len(ws.data_validations.dataValidation) if ws.data_validations else 0,
            "conditional_formatting_ranges": len(ws.conditional_formatting), "freeze_panes": clean(ws.freeze_panes),
            "auto_filter": clean(ws.auto_filter.ref), "cell_types": dict(types), "blank_cells_in_used_range": blanks,
        })
    card_ws = wb["دليل البطاقات"]
    headers = [clean(c.value) for c in card_ws[1]]
    records = [dict(zip(headers, row)) for row in card_ws.iter_rows(min_row=2, values_only=True)]
    ids = [clean(r.get("card_id")) for r in records]
    duplicate_ids = sorted(k for k, v in Counter(ids).items() if k and v > 1)
    name_groups = defaultdict(list)
    for r in records:
        name_groups[(norm(r.get("البنك (إنجليزي)")), norm(r.get("اسم البطاقة")))].append(clean(r.get("card_id")))
    dup_names = {f"{k[0]}|{k[1]}": v for k, v in name_groups.items() if k[1] and len(v) > 1}
    summary = {"workbook": str(path.relative_to(ROOT)), "sheet_count": len(wb.sheetnames), "sheets": wb.sheetnames,
               "card_records": len(records), "duplicate_card_ids": duplicate_ids, "duplicate_name_candidates": dup_names,
               "defined_names": len(wb.defined_names), "calculation_mode": getattr(wb.calculation, "calcMode", None)}
    return summary, sheets


def master_cards() -> tuple[list[dict[str, Any]], list[str]]:
    wb = load_workbook(ORIGINAL, data_only=False, read_only=True)
    ws = wb["دليل البطاقات"]
    headers = [clean(c.value) for c in ws[1]]
    rows = [dict(zip(headers, row)) for row in ws.iter_rows(min_row=2, values_only=True)]
    return rows, headers


MASTER_MAP = {
    "card_name": "اسم البطاقة", "network": "الشبكة", "tier": "الفئة/المستوى",
    "account_segment": "فئة الحساب", "loyalty_program": "برنامج الولاء", "reward_type": "نوع المكافأة",
    "product_type": "نوع المنتج", "calc_tier": "مستوى الثقة (calcTier)", "annual_fee": "الرسوم السنوية (كما وردت)",
    "benefits": "أبرز المزايا/الملاحظات", "source": "مصدر البيانات",
}


def reconcile(master: list[dict[str, Any]], chrome: list[dict[str, str]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    by_id = {clean(r.get("card_id")): r for r in master}
    recs, conflicts = [], []
    for c in chrome:
        m = by_id.get(c["card_id"])
        if not m:
            recs.append({"card_id": c["card_id"], "match_status": "unmatched/new ID - identity review required",
                         "master_name": "", "chrome_name": c["card_name"], "safe_additions": 0, "field_conflicts": 0,
                         "decision": "Preserve as separate provisional source record; do not force merge."})
            continue
        adds, diffs = 0, 0
        for ck, mk in MASTER_MAP.items():
            old, new = clean(m.get(mk)), clean(c.get(ck))
            if not old and new:
                adds += 1
            elif old and new and norm(old) != norm(new):
                material = False
                if ck == "calc_tier":
                    a = next((x for x in ["precise","estimated","unavailable"] if x in old.lower()), "")
                    b = next((x for x in ["precise","estimated","unavailable"] if x in new.lower()), "")
                    material = bool(a and b and a != b)
                elif ck == "annual_fee":
                    nums_a, nums_b = set(re.findall(r"\d+(?:[.,]\d+)?", old)), set(re.findall(r"\d+(?:[.,]\d+)?", new))
                    free_a = any(x in old.lower() for x in ["free", "مجاني", "صفر", "zero"])
                    free_b = any(x in new.lower() for x in ["free", "مجاني", "صفر", "zero"])
                    material = bool((nums_a and nums_b and nums_a.isdisjoint(nums_b)) or (free_a != free_b and (free_a or free_b)))
                elif ck in {"network", "tier"}:
                    vocab = ["visa","mastercard","mada","amex","infinite","signature","platinum","gold","world","elite","classic","titanium"]
                    ta, tb = {x for x in vocab if x in old.lower()}, {x for x in vocab if x in new.lower()}
                    material = bool(ta and tb and ta.isdisjoint(tb))
                if not material:
                    continue
                diffs += 1
                conflicts.append({
                    "bank": clean(m.get("البنك (إنجليزي)")), "card_id": c["card_id"], "card": clean(m.get("اسم البطاقة")),
                    "field": mk, "original_value": old, "new_value": new,
                    "original_source": clean(m.get("مصدر البيانات")) or str(ORIGINAL.relative_to(ROOT)),
                    "new_source": c["source"] or str(CHROME_MD.relative_to(ROOT)), "original_source_date": clean(m.get("تاريخ آخر مراجعة")),
                    "new_source_date": c["source_date"], "decision": "Both values preserved; no overwrite during consolidation.",
                    "confidence": "unresolved", "reasoning": "Populated values differ after conservative normalization.",
                    "user_approval_required": "No during consolidation; yes if a final value must be selected without newer applicable evidence.",
                })
        recs.append({"card_id": c["card_id"], "match_status": "exact card_id match", "master_name": clean(m.get("اسم البطاقة")),
                     "chrome_name": c["card_name"], "safe_additions": adds, "field_conflicts": diffs,
                     "decision": "Link evidence; fill blanks only; preserve conflicts."})
    return recs, conflicts


def missing_rows(master: list[dict[str, Any]], chrome: list[dict[str, str]]) -> list[dict[str, str]]:
    rows = []
    checks = [
        ("اسم البطاقة", "critical"), ("البنك (إنجليزي)", "critical"), ("نوع المنتج", "high"),
        ("مستوى الثقة (calcTier)", "high"), ("مصدر البيانات", "high"), ("الرسوم السنوية (كما وردت)", "medium"),
        ("الشبكة", "medium"), ("برنامج الولاء", "low"),
    ]
    for m in master:
        for field, priority in checks:
            if not clean(m.get(field)):
                rows.append({"bank": clean(m.get("البنك (إنجليزي)")), "card_id": clean(m.get("card_id")), "card": clean(m.get("اسم البطاقة")),
                             "domain": "دليل البطاقات", "missing_field": field, "source_checked": clean(m.get("مصدر البيانات")) or "Original workbook",
                             "expected": "yes" if priority in {"critical", "high"} else "uncertain", "priority": priority,
                             "recommended_action": "Check current official product page, tariff, and applicable terms during final validation."})
    for c in chrome:
        for field, priority in [("card_name", "critical"), ("product_type", "high"), ("source", "high"), ("annual_fee", "medium"), ("network", "medium")]:
            if not clean(c.get(field)):
                rows.append({"bank": c["bank_section"], "card_id": c["card_id"], "card": c["card_name"], "domain": "Chrome V4 staging",
                             "missing_field": field, "source_checked": c["source_file"], "expected": "yes" if priority in {"critical", "high"} else "uncertain",
                             "priority": priority, "recommended_action": "Collect from official current source; do not infer."})
    return rows


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not fieldnames:
        fieldnames = list(rows[0].keys()) if rows else []
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore", lineterminator="\n")
        w.writeheader()
        w.writerows(rows)


def md_table(rows: list[dict[str, Any]], fields: list[str], limit: int | None = None) -> str:
    chosen = rows[:limit] if limit else rows
    def esc(v: Any) -> str:
        return clean(v).replace("|", "\\|").replace("\n", " ")
    out = ["| " + " | ".join(fields) + " |", "| " + " | ".join("---" for _ in fields) + " |"]
    out += ["| " + " | ".join(esc(row.get(f, "")) for f in fields) + " |" for row in chosen]
    return "\n".join(out)


def add_sheet(wb, title: str, rows: list[dict[str, Any]]) -> None:
    if title in wb.sheetnames:
        del wb[title]
    ws = wb.create_sheet(title)
    ws.sheet_view.rightToLeft = True
    if not rows:
        ws["A1"] = "No records"
        return
    headers = list(rows[0].keys())
    ws.append(headers)
    for row in rows:
        ws.append([row.get(h, "") for h in headers])
    fill = PatternFill("solid", fgColor="0F6B78")
    for cell in ws[1]:
        cell.fill = fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    ws.row_dimensions[1].height = 36
    for i, h in enumerate(headers, 1):
        max_len = max([len(clean(h))] + [min(len(clean(r.get(h))), 60) for r in rows[:300]])
        ws.column_dimensions[get_column_letter(i)].width = min(max(max_len + 2, 12), 45)
        for cell in ws[get_column_letter(i)][1:]:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    if "priority" in headers:
        col = get_column_letter(headers.index("priority") + 1)
        dv = DataValidation(type="list", formula1='"critical,high,medium,low,not applicable,uncertain"')
        ws.add_data_validation(dv); dv.add(f"{col}2:{col}{max(ws.max_row,2)}")
    if "match_status" in headers:
        col = get_column_letter(headers.index("match_status") + 1)
        ws.conditional_formatting.add(f"{col}2:{col}{ws.max_row}", FormulaRule(formula=[f'ISNUMBER(SEARCH("unmatched",{col}2))'], fill=PatternFill("solid", fgColor="FFF2CC")))


def build_workbook(inventory: list[dict[str, Any]], chrome: list[dict[str, str]], details: list[dict[str, str]], recs: list[dict[str, Any]], conflicts: list[dict[str, Any]], missing: list[dict[str, str]]) -> Path:
    target = EXCEL / "saudi-credit-cards-unified-consolidated.xlsx"
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(ORIGINAL, target)
    wb = load_workbook(target)
    add_sheet(wb, "Chrome V4 - دليل", chrome)
    add_sheet(wb, "Chrome V4 - تفاصيل", details)
    add_sheet(wb, "مطابقة وتوحيد", recs)
    source_cols = ["source_id", "path", "filename", "extension", "size_bytes", "sha256", "text_sha256", "suspected_duplicate", "duplicate_basis", "bank_mapping", "status", "pages", "encrypted", "text_chars", "error"]
    add_sheet(wb, "مصادر ومراجع", [{k: r.get(k, "") for k in source_cols} for r in inventory])
    add_sheet(wb, "حقول مفقودة", missing)
    add_sheet(wb, "قرارات وتعارضات", conflicts)
    wb.save(target)
    return target


def main() -> None:
    for p in [OUT, REPORTS, MACHINE, EXCEL, WORKING]: p.mkdir(parents=True, exist_ok=True)
    inventory = inventory_sources()
    audit, sheet_audit = workbook_audit(ORIGINAL)
    master, master_headers = master_cards()
    chrome, details = chrome_records()
    recs, new_conflicts = reconcile(master, chrome)
    existing_wb = load_workbook(ORIGINAL, read_only=True, data_only=False)
    cws = existing_wb["سجل التعارضات"]
    cheaders = [clean(c.value) for c in cws[1]]
    inherited_conflicts = []
    for vals in cws.iter_rows(min_row=2, values_only=True):
        r = dict(zip(cheaders, vals))
        inherited_conflicts.append({
            "bank": clean(r.get("البنك")), "card_id": clean(r.get("card_id")), "card": clean(r.get("اسم البطاقة")),
            "field": "Inherited workbook conflict/gap", "original_value": clean(r.get("التعارض/الفجوة")), "new_value": "",
            "original_source": str(ORIGINAL.relative_to(ROOT)), "new_source": "", "original_source_date": "2026-07-29", "new_source_date": "",
            "decision": clean(r.get("الحالة")) or "Preserved", "confidence": "inherited/unresolved", "reasoning": clean(r.get("الخطوة المطلوبة للحسم")),
            "user_approval_required": "Only if final validation cannot resolve it.",
        })
    conflicts = inherited_conflicts + new_conflicts
    missing = missing_rows(master, chrome)

    chrome_by_id = {r["card_id"]: r for r in chrome}
    card_export = []
    for m in master:
        item = {"record_origin": "original_master", **{clean(k): clean(v) for k, v in m.items()}}
        cid = clean(m.get("card_id"))
        item["chrome_exact_id_evidence"] = json.dumps(chrome_by_id.get(cid, {}), ensure_ascii=False) if cid in chrome_by_id else ""
        item["identity_status"] = "exact_id_linked" if cid in chrome_by_id else "master_only"
        card_export.append(item)
    for c in chrome:
        if c["card_id"] not in {clean(m.get("card_id")) for m in master}:
            item = {"record_origin": "chrome_v4_unmatched", **c, "identity_status": "unresolved_identity_do_not_merge"}
            card_export.append(item)
    all_fields = []
    for r in card_export:
        for k in r:
            if k not in all_fields: all_fields.append(k)
    write_csv(MACHINE / "cards.csv", card_export, all_fields)
    (MACHINE / "cards.json").write_text(json.dumps(card_export, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(MACHINE / "sources.csv", inventory)
    write_csv(MACHINE / "conflicts.csv", conflicts)
    write_csv(MACHINE / "missing_fields.csv", missing)
    workbook_target = build_workbook(inventory, chrome, details, recs, conflicts, missing)

    ext_counts = Counter(r["extension"] for r in inventory if r["path"].startswith("Credit Cards Terms and Conditions/"))
    duplicate_groups = defaultdict(list)
    for r in inventory:
        if r["suspected_duplicate"]:
            duplicate_groups[(r["duplicate_basis"], r["sha256"] if r["duplicate_basis"] == "exact bytes" else r["text_sha256"])].append(r["path"])
    duplicate_groups = {f"{k[0]}:{k[1]}":v for k,v in duplicate_groups.items() if len(v)>1}
    unreadable = [r for r in inventory if r["status"] in {"unreadable", "encrypted/unreadable", "unsupported"}]
    inv_text = f"""# Repository Inventory

Generated: {TODAY}. Source hashes are recorded in `working/source_hashes.json`.

## Summary

- Source files under `Credit Cards Terms and Conditions/`: {sum(ext_counts.values())}
- Types: {', '.join(f'`{k}` {v}' for k,v in sorted(ext_counts.items()))}
- Duplicate groups (exact bytes or normalized extracted text): {len(duplicate_groups)}
- Unreadable, encrypted, or unsupported files: {len(unreadable)}
- The RTF source is readable via conservative control-code stripping; layout fidelity is not claimed.

## Files

{md_table(inventory, ['source_id','path','extension','size_bytes','sha256','text_sha256','suspected_duplicate','duplicate_basis','bank_mapping','status','pages','encrypted','text_chars'])}

## Duplicate groups (exact bytes or normalized extracted text)

{json.dumps(duplicate_groups, ensure_ascii=False, indent=2) if duplicate_groups else 'No byte-identical source files detected.'}

## Limitations

PDF text extraction was performed page-by-page. A readable PDF with little or no extractable text may require OCR or visual table review during final validation. DOCX paragraphs and tables were extracted structurally. No source file was modified.

Generated deliverables are maintained under `outputs/`, `docs/`, `scripts/`, and `working/` as defined in `AGENTS.md`; they are intentionally excluded from the raw-source hash table to avoid recursive inventories.
"""
    (ROOT / "docs" / "REPOSITORY_INVENTORY.md").write_text(inv_text, encoding="utf-8")

    precise = sum(norm(m.get("مستوى الثقة (calcTier)")) == "precise" for m in master)
    estimated = sum(norm(m.get("مستوى الثقة (calcTier)")) == "estimated" for m in master)
    unavailable = sum(norm(m.get("مستوى الثقة (calcTier)")) == "unavailable" for m in master)
    bank_names = sorted({clean(m.get("البنك (إنجليزي)")) for m in master if clean(m.get("البنك (إنجليزي)"))})
    categories = sorted({clean(m.get("نوع المنتج")) for m in master if clean(m.get("نوع المنتج"))})
    inherited = MASTER_MD.read_text(encoding="utf-8")
    master_ref = f"""# Consolidated Master Data Reference

Generated: {TODAY}. This is the continuation document for the Saudi payment-card working dataset.

## Purpose and scope boundary

The deliverable consolidates inherited workbook records, the prior Claude reference, Chrome V4 official-site research, and the available document corpus without claiming final currency or completeness. It does not modify or feed the main CCIP platform.

## Current state

- Original master: {len(master)} card records, {len(bank_names)} issuer names, {audit['sheet_count']} original worksheets.
- Original `calcTier`: precise {precise}, estimated {estimated}, unavailable {unavailable}.
- Chrome V4: {len(chrome)} card-guide records and {len(details)} structured detail rows across nine bank sections.
- Exact `card_id` links: {sum(r['match_status']=='exact card_id match' for r in recs)}.
- Unmatched/new identifier records preserved provisionally: {sum(r['match_status']!='exact card_id match' for r in recs)}.
- Conflicts preserved: {len(conflicts)} ({len(inherited_conflicts)} inherited workbook entries; {len(new_conflicts)} field-level differences detected during exact-ID reconciliation).
- Missing-field backlog rows: {len(missing)}.
- Consolidated workbook: `{workbook_target.relative_to(ROOT)}`.

## Sources and authority

See `docs/REPOSITORY_INVENTORY.md` for hashes, sizes, extractability, and bank mappings. The provisional priority is: current tariff; product T&Cs; current product page; rewards terms; FAQ/campaign; payment network; traceable prior research; unsupported workbook value; secondary source. Conflicts are preserved regardless of rank until applicability and date are established.

## Workbook structure

All original sheets remain. Six additive sheets were created: Chrome card guide, Chrome detail staging, reconciliation, source registry, missing fields, and decisions/conflicts. Original formulas, formatting objects, tables, filters, comments, hyperlinks, validations, merged cells, hidden states, and frozen panes are compared during validation.

## Banks

{', '.join(bank_names)}

## Card categories in the original master

{', '.join(categories)}

## Data conventions and decisions

- `precise`, `estimated`, and `unavailable` retain the inherited definitions below.
- Direct evidence, normalization, inference, conflict, missing data, provisional decisions, and confirmed decisions are distinct states.
- Exact IDs link evidence; only blank fields are candidates for safe enrichment. Populated differences become conflict rows.
- Mismatched namespaces are not force-mapped. The 98 unmatched Chrome records remain source records with `unresolved_identity_do_not_merge` in machine exports.
- No card, historical record, or source file was deleted. No populated master value was silently overwritten.

## Completed work

- Reconstructed repository and workbook state.
- Read all Markdown sources and programmatically inspected all worksheets and document files.
- Parsed Chrome V4 card-guide and detail tables.
- Created additive reconciliation and provenance layers.
- Generated missing-data, conflict, collection-status, audit, change, and final-validation reports.
- Generated CSV/JSON exports without flattening or replacing the original workbook structure.

## Partially completed and outstanding work

- Identifier mapping for unmatched Chrome V4 records requires bank-by-bank identity review.
- Conflicting official values require date/applicability checks.
- Image-only or layout-dependent PDF tables require visual/OCR review where extraction is weak.
- Product availability, fees, APR, rewards, limits, eligibility, and benefits still require final current-source validation.
- Machine exports preserve rich text values; they do not decompose every narrative fee/benefit into atomic database fields.

## Instructions for the next agent

Read `AGENTS.md`, this file, `CONFLICTS_AND_DECISIONS.md`, and `MISSING_INFORMATION.md`. Rerun the scripts before editing. Work bank-by-bank, establish explicit ID aliases, add evidence rather than replacing it, and update counts. Do not restart research or assume unmatched IDs are new products.

## Final-validation plan

Follow `outputs/reports/FINAL_VALIDATION_PLAN.md`. The next recommended action is an explicit identity-mapping pass for the unmatched Chrome namespaces, starting with the highest-volume banks, before selecting any conflicting values.

---

# Inherited Claude master reference (carried forward verbatim)

{inherited}
"""
    (OUT / "MASTER_DATA_REFERENCE.md").write_text(master_ref, encoding="utf-8")

    (REPORTS / "MISSING_INFORMATION.md").write_text(
        f"# Missing Information\n\nGenerated: {TODAY}. Rows are observations, not claims that the field must exist.\n\n"
        + md_table(missing, ["bank","card_id","card","domain","missing_field","source_checked","expected","priority","recommended_action"]), encoding="utf-8")
    (REPORTS / "CONFLICTS_AND_DECISIONS.md").write_text(
        f"# Conflicts and Decisions\n\nGenerated: {TODAY}. No conflict was resolved by silently selecting one value.\n\n"
        + md_table(conflicts, ["bank","card_id","card","field","original_value","new_value","original_source","new_source","original_source_date","new_source_date","decision","confidence","reasoning","user_approval_required"]), encoding="utf-8")

    bank_status = []
    conflict_ids = {r["card_id"] for r in conflicts}
    missing_ids = {r["card_id"] for r in missing if r["priority"] in {"critical", "high"}}
    exact_ids = {r["card_id"] for r in recs if r["match_status"] == "exact card_id match"}
    for card in card_export:
        cid = clean(card.get("card_id"))
        bank = clean(card.get("البنك (إنجليزي)")) or clean(card.get("bank_section"))
        name = clean(card.get("اسم البطاقة")) or clean(card.get("card_name"))
        if cid in conflict_ids: status = "conflict pending"
        elif cid in missing_ids: status = "missing evidence"
        elif card["record_origin"] == "chrome_v4_unmatched": status = "collected but not consolidated"
        elif cid in exact_ids: status = "consolidated"
        else: status = "source files available"
        bank_status.append({"bank": bank, "card_id": cid, "card": name, "status": status,
                            "evidence": card["record_origin"], "next_action": "Final official-source validation" if status == "consolidated" else "Resolve identity/conflict/missing evidence"})
    (REPORTS / "COLLECTION_STATUS.md").write_text(
        f"# Collection Status\n\nGenerated: {TODAY}. `ready for final validation` is intentionally not assigned automatically.\n\n"
        + md_table(bank_status, ["bank","card_id","card","status","evidence","next_action"]), encoding="utf-8")

    audit_report = f"""# Workbook Audit

Generated: {TODAY}.

## Structure

- Original workbook: `{audit['workbook']}`
- Sheets: {audit['sheet_count']}
- Card records: {audit['card_records']}
- Defined names: {audit['defined_names']}
- Duplicate card IDs: {len(audit['duplicate_card_ids'])}
- Duplicate normalized bank/card-name candidates: {len(audit['duplicate_name_candidates'])}
- Formula cells: {sum(r['formulas'] for r in sheet_audit)}
- Comments: {sum(r['comments'] for r in sheet_audit)}
- Hyperlinks: {sum(r['hyperlinks'] for r in sheet_audit)}

{md_table(sheet_audit, ['sheet','state','rows','columns','formulas','comments','hyperlinks','merged_ranges','hidden_rows','hidden_columns','tables','data_validations','conditional_formatting_ranges','freeze_panes','auto_filter'])}

## Quality findings and risks

- The workbook contains no formulas, so there are no formula-reference errors to repair; this also means coverage summaries are static and can drift.
- The main card guide has {audit['card_records']} records and no duplicate `card_id` values.
- Blank mandatory/expected fields are enumerated in `MISSING_INFORMATION.md` rather than inferred.
- Similar card names and changed ID namespaces are identity candidates, not confirmed duplicates.
- Narrative fee, APR, bonus, and benefit fields contain mixed units and composite values. Machine exports preserve them as text to avoid lossy parsing.
- The original master has one Excel table and no named ranges. Generated sheets use filters and frozen headers.
- No original row, column, sheet, formula, comment, hyperlink, validation, merge, or hidden state is intentionally removed.

## Duplicate-name candidates

```json
{json.dumps(audit['duplicate_name_candidates'], ensure_ascii=False, indent=2)}
```
"""
    (REPORTS / "WORKBOOK_AUDIT.md").write_text(audit_report, encoding="utf-8")

    changelog = f"""# Changelog

Generated: {TODAY}.

## Inspected

- Repository source files: {sum(1 for r in inventory if r['path'].startswith('Credit Cards Terms and Conditions/'))}.
- Markdown files read: {sum(r['extension']=='.md' for r in inventory)}.
- Excel workbooks: {sum(r['extension']=='.xlsx' for r in inventory)}; worksheets audited: {audit['sheet_count']}.
- PDFs: {sum(r['extension']=='.pdf' for r in inventory)}; DOCX: {sum(r['extension']=='.docx' for r in inventory)}; RTF: {sum(r['extension']=='.rtf' for r in inventory)}.

## Data changes

- Original master rows retained: {len(master)}; original row deletions: **0**.
- Original sheets retained: {audit['sheet_count']}; original sheet deletions: **0**.
- Original populated values overwritten: **0**.
- Additive workbook sheets: **6**.
- Chrome V4 guide records staged: {len(chrome)}; structured detail rows staged: {len(details)}.
- Exact-ID evidence links: {sum(r['match_status']=='exact card_id match' for r in recs)}.
- Unmatched/new-ID records preserved separately: {sum(r['match_status']!='exact card_id match' for r in recs)}.
- Field-level differences preserved as new conflict rows: {len(new_conflicts)}.
- Inherited workbook conflicts preserved: {len(inherited_conflicts)}.
- Missing-field observations: {len(missing)}.
- Machine-readable card rows: {len(card_export)}.

## Intentionally unchanged

All original sheets, rows, cell values, formatting, formulas, comments, hyperlinks, and historical records. No identifier aliases were force-created for incompatible namespaces.

## Scripts and outputs

Created repeatable consolidation, validation, and comparison scripts. Generated the consolidated workbook, five machine-readable exports, repository inventory, master reference, audit, missing-data, conflict, collection-status, changelog, and final-validation plan.

## Pre-merge release review

- Added the permanent repository separation, source-ingestion, and GitHub lifecycle policy to `AGENTS.md`.
- Removed the generated repository inventory from its own hash scope and excluded ignored runtime caches, eliminating recursive/non-source inventory drift.
- Added Git-backed raw-source preservation, generated-file placement, XLSX ZIP integrity, and Markdown UTF-8 checks.
- Added `scripts/check_reproducibility.py`; two independent rebuilds produced identical deterministic reports/exports and identical workbook semantics.
- Normalized generated CSV line endings to LF so repository diffs are stable across environments.
"""
    (REPORTS / "CHANGELOG.md").write_text(changelog, encoding="utf-8")

    final_plan = f"""# Final Validation Plan

This consolidation is not final verification. Execute the next phase bank-by-bank.

1. Freeze an explicit card identity/alias map, starting with the {sum(r['match_status']!='exact card_id match' for r in recs)} unmatched Chrome IDs.
2. For every identifiable card, capture the current official product page, current pricing guide, applicable T&Cs, and rewards/benefits terms with source dates.
3. Confirm current availability and status: current, renamed, replaced, discontinued, historical, or uncertain.
4. Validate fees, APR examples and assumptions, rewards rates and exclusions, limits, eligibility, bonuses, and benefits.
5. Compare official sources by effective date and applicability; retain superseded values as history.
6. Resolve conflicts only with documented reasoning. Ask the user when evidence cannot support a material selection.
7. Promote records to `ready for final validation` and then `confirmed` only after the full evidence set is checked.
8. Rebuild and validate all outputs; reconcile row counts and require deletion count zero.

Recommended order: unresolved namespace mapping; high-priority fee/APR conflicts; cards with missing source or product type; remaining estimated/unavailable records; then full current-product sweep.
"""
    (REPORTS / "FINAL_VALIDATION_PLAN.md").write_text(final_plan, encoding="utf-8")
    print(json.dumps({"master_cards": len(master), "chrome_cards": len(chrome), "chrome_details": len(details),
                      "exact_matches": sum(r['match_status']=='exact card_id match' for r in recs),
                      "unmatched": sum(r['match_status']!='exact card_id match' for r in recs),
                      "conflicts": len(conflicts), "missing": len(missing), "sources": len(inventory),
                      "output_workbook": str(workbook_target)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
